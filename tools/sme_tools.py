from typing import List
from langchain.tools import tool
from fpdf import FPDF
import os
import io
import subprocess
import tempfile
from pathlib import Path
from serpapi.google_search import GoogleSearch
from email.message import EmailMessage
import smtplib
from dotenv import load_dotenv
import shutil
load_dotenv()
# @tool
# def search(query: str) -> str:
#     """Search for information."""
#     return f"Results for: {query}"

# @tool
# def get_weather(location: str) -> str:
#     """Get weather information for a location."""
#     return f"Weather in {location}: Sunny, 72°F"

@tool
def create_docx_report(title: str, content: str, image_paths:List[str]=None) -> str:
    """
    Creates a simple Microsoft Word (.docx) document with a title and content.

    Args:
        title: The main heading of the document.
        content: The body text of the document.

    Returns:
        The absolute path to the saved .docx file.
    """
    # Sanitize the title to create a valid filename
    import os
    from docx import Document
    from fpdf import FPDF
    from docx.shared import Inches
    OUTPUT_DIR = "outputs/generated_reports"
    safe_filename = title.replace(' ', '_').replace('/', '_')
    file_path = os.path.join(OUTPUT_DIR, f"{safe_filename}.docx")

    doc = Document()
    doc.add_heading(title, level=0)  # Add a main title
    
    # Add paragraphs, preserving line breaks from the original content
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    
    if image_paths:
        doc.add_paragraph("\n--- Images ---")
        for img_path in image_paths:
            img_path = os.path.join("outputs/images/", img_path)
            if os.path.exists(img_path):
                try:
                    with open(img_path, 'rb') as img_file:
                        image_bytes = img_file.read()
                    img_stream = io.BytesIO(image_bytes)
                    doc.add_picture(img_stream)
                    doc.add_paragraph(f"Image: {os.path.basename(img_path)}", style='Caption')
                except Exception as e:
                    print(f"Warning: Could not add image {img_path}. Error: {e}")
            else:
                print(f"Warning: Image path not found: {img_path}")
    doc.save(file_path)
    # print(f"DOCX report saved to: {file_path}")
    return f"DOCX report has been saved to {os.path.abspath(file_path)}"

class PDF(FPDF):
    """
    Custom PDF class to handle headers and footers automatically.
    """
    def header(self):
        self.set_font('Arial', 'B', 12)
        # Move to the right
        self.cell(80)
        # Title
        self.cell(30, 10, self.title, 0, 0, 'C')
        # Line break
        self.ln(20)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        # Page number
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

@tool
def create_pdf_report(title: str, content: str) -> str:
    """
    Creates a simple PDF document with a title and content using FPDF2.

    Args:
        title: The title of the document.
        content: The body text of the document.

    Returns:
        The absolute path to the saved .pdf file.
    """
    import os
    from fpdf import FPDF
    OUTPUT_DIR = "outputs/generated_reports"
    safe_filename = title.replace(' ', '_').replace('/', '_')
    file_path = os.path.join(OUTPUT_DIR, f"{safe_filename}.pdf")

    pdf = FPDF()
    pdf.set_title(title)
    pdf.add_page()

    font_path = os.path.join(os.path.dirname(__file__), "../fonts", "Arial Unicode.ttf")
    pdf.add_font("ArialUnicode", "", font_path, uni=True)
    pdf.set_font("ArialUnicode", "", 12)

    pdf.multi_cell(0, 10, content)
    pdf.output(file_path)
    # print(f"PDF report saved to: {file_path}")
    return f"PDF report has been saved to {os.path.abspath(file_path)}"

@tool
def compile_latex_to_pdf(latex_code: str, title: str, image_paths: List[str] = None) -> str:
    """
    Compiles LaTeX code into a PDF file using pdflatex.

    Args:
        latex_code (str): The LaTeX document code as a string.
        output_dir (str): Directory to save the compiled PDF. Defaults to current directory.

    Returns:
        str: Absolute path to the generated PDF file.

    Raises:
        RuntimeError: If compilation fails or pdflatex is not found.
    """
    # Ensure output directory exists
    output_dir = "outputs/generated_reports"
    os.makedirs(output_dir, exist_ok=True)

    # Create a temporary working directory for compilation
    with tempfile.TemporaryDirectory() as tmpdir:
        tex_path = Path(tmpdir) / "document.tex"
        pdf_path = Path(tmpdir) / "document.pdf"

        # Write LaTeX code to .tex file
        tex_path.write_text(latex_code, encoding="utf-8")
        if image_paths:
            for img_path in image_paths:
                # img_path = os.path.join("outputs/images/", img_path)
                if os.path.exists(img_path):
                    # Copy the image file into the temp directory
                    # so \includegraphics{filename.png} can find it.
                    shutil.copy(img_path, tmpdir)
                else:
                    print(f"Warning: Image path not found, skipping: {img_path}")
        # Run pdflatex (suppressing console output)
        try:
            subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", tex_path.name],
                cwd=tmpdir,
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
        except subprocess.CalledProcessError as e:
            error_log = (Path(tmpdir) / "document.log").read_text(encoding="utf-8", errors="ignore")
            raise RuntimeError(f"LaTeX compilation failed:\n{error_log}") from e
        except FileNotFoundError:
            raise RuntimeError("pdflatex not found. Please install TeX Live or MiKTeX.")

        # Move compiled PDF to desired output_dir
        safe_filename = title.replace(' ', '_').replace('/', '_')
        if safe_filename.lower().endswith('.pdf') == 0:
            safe_filename += ".pdf"
        final_path = Path(output_dir) / safe_filename
        final_path.write_bytes(pdf_path.read_bytes())

    return f"PDF successfully created at: {final_path.resolve()}"

@tool
def get_image_links(query, limit=10, api_key="40bdda4a2a020a149f801ff52409bd172cb378dbb26c9b772ca14e4fc8f28972"):
    """
    Fetch image links from Google Images using SerpAPI.
    """
    search = GoogleSearch({
        "q": query,
        "tbm": "isch",  # image search
        "num": limit,
        "api_key": api_key
    })
    results = search.get_dict()
    return [img["original"] for img in results.get("images_results", [])]

@tool
def send_email(to, subject, body, attachments=None):
    """
    Send an email via Gmail SMTP with optional attachments.

    Args:
        sender (str): Your Gmail address
        app_password (str): Gmail app password
        to (str | list): Recipient email(s)
        subject (str): Email subject
        body (str): Body text
        attachments (list[str], optional): Paths to attachments
    """
    msg = EmailMessage()
    msg["From"] = "cosmiccompass.lma@gmail.com"
    msg["To"] = ", ".join(to) if isinstance(to, list) else to
    msg["Subject"] = subject
    msg.set_content(body)
    # Add attachments if any
    if attachments:
        for path in attachments:
            name = os.path.basename(path)
            dirname = os.path.dirname(path)
            name = name.replace(" ", "_")
            path = os.path.join(dirname, name)
            with open(path, "rb") as f:
                data = f.read()
                name = os.path.basename(path)
            msg.add_attachment(data, maintype="application", subtype="octet-stream", filename=name)

    # Connect and send
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login("cosmiccompass.lma@gmail.com", os.getenv("SMTP_API"))
        smtp.send_message(msg)
    return "Email sent successfully!"
